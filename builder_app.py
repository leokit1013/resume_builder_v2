import streamlit as st
import os
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
from dotenv import load_dotenv
from langdetect import detect
import easyocr
import numpy as np
import fitz
from io import BytesIO
from fpdf import FPDF
from docx import Document
from fpdf import FPDF
import unicodedata
import re
from datetime import date
from streamlit_webrtc import webrtc_streamer, AudioProcessorBase
import av
import tempfile
# from weasyprint import HTML


load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in .env file or environment variables.")

genai.configure(api_key=GOOGLE_API_KEY)

model = genai.GenerativeModel(
    'gemini-1.5-flash-8b-001',
    generation_config={
        "temperature": 0.7,
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }
)


# Define Gemini Pro Audio model
audio_model = genai.GenerativeModel("models/gemini-1.5-pro-latest")

class AudioProcessor(AudioProcessorBase):
    def __init__(self):
        self.frames = []

    def recv(self, frame: av.AudioFrame) -> av.AudioFrame:
        self.frames.append(frame)
        return frame

def transcribe_with_gemini_from_audio(frames):
    if not frames:
        return "No audio frames captured."

    try:
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
            # Open a writable WAV container
            container = av.open(temp_audio.name, mode='w', format='wav')
            stream = container.add_stream("pcm_s16le", rate=frames[0].sample_rate)
            stream.layout = frames[0].layout  # Set layout only (not channels directly)

            for frame in frames:
                container.mux(frame)
            container.close()

            # Read the WAV file
            with open(temp_audio.name, "rb") as audio_file:
                audio_data = audio_file.read()

        # Gemini transcription
        try:
            response = audio_model.generate_content([
                genai.content.FileData(mime_type="audio/wav", data=audio_data)
            ])
            return response.text

        except Exception as e:
            st.error(f"Gemini transcription failed: {e}")
            return ""

    except Exception as e:
        st.error(f"Audio processing failed: {e}")
        return ""


def detect_language(text):
    return detect(text)

def read_image(file, lang):
    image = Image.open(file)
    image_np = np.array(image)
    reader = easyocr.Reader([lang, 'en'], gpu=False)
    result = reader.readtext(image_np, detail=0)
    return ' '.join(result)




def clean_text(text):
    # Normalize Unicode and remove unsupported characters for FPDF
    return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")

# ---------- Markdown to HTML Converter ----------
def markdown_to_html(text):
    html = text
    html = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", html)  # Bold
    html = re.sub(r"\*(.*?)\*", r"<em>\1</em>", html)              # Italic
    html = html.replace("\n", "<br>")                              # Newlines to <br>
    return f"<div style='font-family:Arial; font-size:15px;'>{html}</div>"

# ---------- Save as DOCX ----------
def save_as_docx(text, filename="resume.docx"):
    doc = Document()

    lines = text.split("\n")
    for line in lines:
        paragraph = doc.add_paragraph()

        # Bold: **text**
        line = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line)
        # Italic: *text*
        line = re.sub(r"\*(.*?)\*", r"<i>\1</i>", line)

        # Split line by HTML tags for bold and italic
        tokens = re.split(r"(<b>.*?</b>|<i>.*?</i>)", line)
        for token in tokens:
            if token.startswith("<b>") and token.endswith("</b>"):
                run = paragraph.add_run(token[3:-4])
                run.bold = True
            elif token.startswith("<i>") and token.endswith("</i>"):
                run = paragraph.add_run(token[3:-4])
                run.italic = True
            else:
                paragraph.add_run(token)
    doc.save(filename)
    return filename

# ---------- Save as PDF ----------
def save_as_pdf(text, filename="resume.pdf"):
    html = markdown_to_html(text)

    output_path = f"{filename}"
    HTML(string=html).write_pdf(output_path)
    return output_path



def extract_text_from_file(file):
    try:
        if file.name.endswith(".pdf"):
            text = ""
            pdf_bytes = file.read()
            if not pdf_bytes:
                st.error("Error: Uploaded PDF is empty.")
                return None
            try:
                reader = PdfReader(BytesIO(pdf_bytes))
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n\n"
            except:
                pass
            if not text:
                try:
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    for page in doc:
                        t = page.get_text("text")
                        if t:
                            text += t + "\n\n"
                except:
                    pass
            if not text:
                try:
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    for page in doc:
                        pix = page.get_pixmap()
                        img_bytes = pix.tobytes("png")
                        lang = detect_language(read_image(BytesIO(img_bytes), 'en'))
                        text += read_image(BytesIO(img_bytes), lang) + "\n\n"
                except Exception as e:
                    st.error(f"OCR failed: {e}")
                    return None
            return text

        elif file.name.endswith(".txt"):
            return file.read().decode("utf-8")

        elif file.name.endswith(".docx"):
            doc = Document(file)
            return "\n".join(p.text for p in doc.paragraphs)

        elif file.name.endswith((".jpg", ".jpeg", ".png")):
            temp_text = read_image(file, 'en')
            lang = detect_language(temp_text)
            return read_image(file, lang)

        else:
            st.error("Unsupported file type.")
            return None

    except Exception as e:
        st.error(f"Text extraction failed: {e}")
        return None

def summarize_with_gemini(text, prompt_template):
    try:
        prompt = prompt_template.format(text=text)
        response = model.generate_content(prompt)
        return response.text if response.text else None
    except Exception as e:
        st.error(f"Gemini API Error: {e}")
        return None

def build_user_input(data):
    lines = []

    # Personal Info
    pi = data.get("personal_info", {})
    if pi:
        lines.append(f"Name: {pi.get('name', '')}")
        lines.append(f"Email: {pi.get('email', '')}")
        lines.append(f"Phone: {pi.get('phone', '')}")
        lines.append(f"Summary: {pi.get('summary', '')}")
    
    # Experience
    if data.get("experience"):
        lines.append("\nExperience:")
        for exp in data["experience"]:
            if isinstance(exp, dict):
                lines.append(f"{exp.get('job_title', 'N/A')} at {exp.get('company', 'N/A')} ({exp.get('start', 'N/A')} to {exp.get('end', 'N/A')})")
                lines.append(f"Description: {exp.get('description', '')}")
            else:
                lines.append(str(exp))  # fallback if not dict
    
    # Education
    if data.get("education"):
        lines.append("\nEducation:")
        for edu in data["education"]:
            if isinstance(edu, dict):
                lines.append(f"{edu.get('degree', 'N/A')} from {edu.get('school', 'N/A')} ({edu.get('year', 'N/A')})")
            else:
                lines.append(str(edu))

    # Skills
    if data.get("skills"):
        if isinstance(data["skills"], list):
            lines.append(f"\nSkills: {', '.join(map(str, data['skills']))}")
        else:
            lines.append(f"\nSkills: {str(data['skills'])}")

    # Certifications
    if data.get("certifications"):
        lines.append("\nCertifications:")
        for cert in data["certifications"]:
            if isinstance(cert, dict):
                lines.append(f"{cert.get('name', 'Unknown')} issued by {cert.get('org', 'Unknown')}")
            else:
                lines.append(str(cert))  # fallback if it's just a string

    return "\n".join([l for l in lines if l.strip()])


def voice_input_section(label: str, key: str):
    st.markdown(f"🎤 **{label}**")

    webrtc_ctx = webrtc_streamer(
        key=f"rec-{key}",
        audio_receiver_size=4096,
        media_stream_constraints={"audio": True, "video": False},
        async_processing=True,
        audio_processor_factory=AudioProcessor,
    )

    transcript = st.empty()

    # Automatically transcribe once user stops recording
    if webrtc_ctx and not webrtc_ctx.state.playing and hasattr(webrtc_ctx.audio_processor, "frames"):
        frames = webrtc_ctx.audio_processor.frames
        if frames:
            with st.spinner("Transcribing with Gemini..."):
                result = transcribe_with_gemini_from_audio(frames)
                st.session_state.resume_data["skills_from_voice"] = result
            transcript.text_area("Transcript", value=result, key=f"text-{key}")
            return result
        else:
            st.warning("No audio recorded.")

    return ""

def display_edit_resume():
    """Display editable resume fields from session_state."""
    if "editable_resume" not in st.session_state:
        st.warning("No resume to edit yet. Please generate one first.")
        return

    resume = st.session_state.editable_resume

    st.subheader("👤 Personal Information")
    resume["personal_info"]["name"] = st.text_input("Full Name", resume["personal_info"].get("name", ""))
    resume["personal_info"]["email"] = st.text_input("Email", resume["personal_info"].get("email", ""))
    resume["personal_info"]["phone"] = st.text_input("Phone", resume["personal_info"].get("phone", ""))
    resume["personal_info"]["summary"] = st.text_area("Summary", resume["personal_info"].get("summary", ""))

    st.subheader("🛠 Skills")
    skills_text = st.text_area("Skills (comma-separated)", ", ".join(resume.get("skills", [])))
    resume["skills"] = [s.strip() for s in skills_text.split(",") if s.strip()]

    st.subheader("💼 Experience")
    for i, exp in enumerate(resume.get("experience", [])):
        exp["job_title"] = st.text_input(f"Job Title {i+1}", exp.get("job_title", ""))
        exp["company"] = st.text_input(f"Company {i+1}", exp.get("company", ""))
        col1, col2 = st.columns(2)
        with col1:
            exp["start_date"] = st.text_input(f"Start Date {i+1}", exp.get("start_date", ""))
        with col2:
            exp["end_date"] = st.text_input(f"End Date {i+1}", exp.get("end_date", ""))
        exp["description"] = st.text_area(f"Description {i+1}", exp.get("description", ""))

    if st.button("💾 Save Edits"):
        st.session_state.editable_resume = resume
        st.success("Edits saved!")

    if st.button("⬇️ Download Updated Resume"):
        final_resume_text = build_user_input(resume)  # Convert back to text/markdown
        docx_path = save_as_docx(final_resume_text, "updated_resume.docx")
        st.download_button("📄 Download as DOCX", data=open(docx_path, "rb"), file_name="updated_resume.docx")

# --- Convert dict back to markdown ---
def build_resume_markdown(resume_data: dict) -> str:
    """Convert structured resume dict back into markdown format."""
    md = ""

    # Profile photo at the top
    if "profile_photo" in st.session_state:
        md += f'<img src="data:image/png;base64,{st.session_state["profile_photo_base64"]}" width="120" style="float:right;margin-left:20px;border-radius:10px;">\n'

    md += f"## {resume_data['personal_info'].get('name', '')}\n"
    md += f"**Email:** {resume_data['personal_info'].get('email', '')}  \n"
    md += f"**Phone:** {resume_data['personal_info'].get('phone', '')}  \n\n"
    md += f"### Summary\n{resume_data['personal_info'].get('summary', '')}\n\n"

    if resume_data.get("skills"):
        md += "### Skills\n" + ", ".join(resume_data["skills"]) + "\n\n"

    if resume_data.get("experience"):
        md += "### Experience\n"
        for exp in resume_data["experience"]:
            md += f"**{exp.get('job_title', '')}** at {exp.get('company', '')}  \n"
            md += f"{exp.get('start_date', '')} – {exp.get('end_date', '')}  \n"
            md += f"{exp.get('description', '')}\n\n"

    if resume_data.get("education"):
        md += "### Education\n"
        for edu in resume_data["education"]:
            md += f"**{edu.get('degree', '')}**, {edu.get('institution', '')} ({edu.get('year', '')})\n"

    return md


def parse_generated_resume(resume_text: str) -> dict:
    """
    Parses the generated resume text into a structured dictionary.
    Works best if resume sections have headings like 'Summary', 'Skills', 'Experience', etc.
    """
    parsed_resume = {
        "personal_info": {
            "name": "",
            "email": "",
            "phone": "",
            "summary": ""
        },
        "skills": [],
        "experience": [],
        "education": []
    }

    # --- Extract Personal Info ---
    # These patterns may need adjusting based on Gemini output format
    name_match = re.search(r"Name\s*:\s*(.+)", resume_text, re.IGNORECASE)
    email_match = re.search(r"Email\s*:\s*(.+)", resume_text, re.IGNORECASE)
    phone_match = re.search(r"Phone\s*:\s*(.+)", resume_text, re.IGNORECASE)

    if name_match:
        parsed_resume["personal_info"]["name"] = name_match.group(1).strip()
    if email_match:
        parsed_resume["personal_info"]["email"] = email_match.group(1).strip()
    if phone_match:
        parsed_resume["personal_info"]["phone"] = phone_match.group(1).strip()

    # --- Extract Summary ---
    summary_match = re.search(r"Summary\s*[:\-]?\s*(.*?)(?:\n\n|\Z)", resume_text, re.IGNORECASE | re.DOTALL)
    if summary_match:
        parsed_resume["personal_info"]["summary"] = summary_match.group(1).strip()

    # --- Extract Skills ---
    skills_match = re.search(r"Skills\s*[:\-]?\s*(.*?)(?:\n\n|\Z)", resume_text, re.IGNORECASE | re.DOTALL)
    if skills_match:
        skills_text = skills_match.group(1).strip()
        parsed_resume["skills"] = [s.strip() for s in re.split(r",|\n", skills_text) if s.strip()]

    # --- Extract Experience ---
    exp_section = re.search(r"Experience\s*[:\-]?\s*(.*?)(?:\n(?:Education|Skills|Certifications)|\Z)",
                            resume_text, re.IGNORECASE | re.DOTALL)
    if exp_section:
        exp_entries = re.split(r"\n\s*\n", exp_section.group(1).strip())  # split by double newlines
        for entry in exp_entries:
            lines = entry.strip().split("\n")
            if not lines:
                continue
            parsed_resume["experience"].append({
                "job_title": lines[0] if len(lines) > 0 else "",
                "company": lines[1] if len(lines) > 1 else "",
                "start_date": "",
                "end_date": "",
                "description": "\n".join(lines[2:]) if len(lines) > 2 else ""
            })

    # --- Extract Education ---
    edu_section = re.search(r"Education\s*[:\-]?\s*(.*?)(?:\n(?:Experience|Skills|Certifications)|\Z)",
                            resume_text, re.IGNORECASE | re.DOTALL)
    if edu_section:
        edu_entries = re.split(r"\n\s*\n", edu_section.group(1).strip())
        for entry in edu_entries:
            lines = entry.strip().split("\n")
            if not lines:
                continue
            parsed_resume["education"].append({
                "degree": lines[0] if len(lines) > 0 else "",
                "institution": lines[1] if len(lines) > 1 else "",
                "year": lines[2] if len(lines) > 2 else ""
            })

    return parsed_resume

def main():
    st.set_page_config(page_title="Resume Builder & Enhancer", layout="wide")
    st.title("🧠 Smart Resume Generator & Improver")

    tab1, tab2 = st.tabs(["📝 Create Resume", "🔧 Improve Resume"])

    with tab1:
            
        # Configure page
        # Inject CSS for modern UI
                
        st.markdown("""
        <style>
            /* Center page and style container */
            .block-container {
                padding-top: 1rem;
                padding-bottom: 1rem;
                max-width: 1200px;
                margin: auto;
            }
            /* Form card styling */
            .form-card {
                background: white;
                padding: 1.5rem;
                border-radius: 12px;
                box-shadow: 0px 4px 12px rgba(0,0,0,0.05);
            }
            /* Photo upload card */
            .photo-card {
                background: #f8faff;
                border: 2px dashed #3b82f6;
                border-radius: 12px;
                text-align: center;
                padding: 2rem;
            }
            .photo-card img {
                max-width: 100%;
                border-radius: 10px;
            }
        </style>
        """, unsafe_allow_html=True)



        # Initialize session state
        if "resume_data" not in st.session_state:
            st.session_state.resume_data = {
                "personal_info": {},
                "experience": [],
                "education": [],
                "skills": "",
                "certifications": ""
            }

        st.title("📄 AI Resume Builder")

        # Horizontal tabs for navigation
        tabs = st.tabs([
            "Personal Info", "Experience", "Education", 
            "Skills", "Certifications", "Preview Resume"
        ])

        # ==== PERSONAL INFO TAB ====
        with tabs[0]:
            st.subheader("👤 Step 1: Personal Info")
            col_left, col_right = st.columns([2, 1])

            with col_left:
                with st.form("personal_form", clear_on_submit=False):
                    c1, c2 = st.columns(2)
                    with c1:
                        first_name = st.text_input("First Name")
                    with c2:
                        last_name = st.text_input("Last Name")

                    email = st.text_input("Email Address")
                    phone = st.text_input("Phone Number")
                    location = st.text_input("Location (City, State, Country)")

                    c3, c4 = st.columns(2)
                    with c3:
                        linkedin = st.text_input("LinkedIn Profile (optional)")
                    with c4:
                        website = st.text_input("Personal Website (optional)")

                    summary = st.text_area("Summary (optional)")
                    voice_summary = voice_input_section("Summary via Voice", "voice_summary")

                    if st.form_submit_button("Save Personal Info"):
                        full_summary = summary + "\n" + voice_summary if voice_summary else summary
                        st.session_state.resume_data["personal_info"] = {
                            "name": f"{first_name} {last_name}",
                            "email": email,
                            "phone": phone,
                            "location": location,
                            "linkedin": linkedin,
                            "website": website,
                            "summary": full_summary
                        }
                        st.success("✅ Saved Personal Info")

            with col_right:
                uploaded_photo = st.file_uploader("Upload Photo", type=["jpg", "jpeg", "png"])
                if uploaded_photo:
                    st.session_state["profile_photo"] = uploaded_photo.getvalue()

                    # Save photo into resume_data
                    if "resume_data" not in st.session_state:
                        st.session_state.resume_data = {}
                    st.session_state.resume_data["photo"] = st.session_state["profile_photo"]

                if "profile_photo" in st.session_state:
                    st.image(st.session_state["profile_photo"], width=150)
        # ==== EXPERIENCE TAB ====
        with tabs[1]:
            st.subheader("💼 Experience")
            col_left, col_right = st.columns([2, 1])
            with col_left:
                with st.form("experience_form"):
                    job_title = st.text_input("Job Title")
                    company = st.text_input("Company Name")
                    start_date = st.date_input("Start Date", value=date(2020, 1, 1))
                    end_date = st.date_input("End Date", value=date.today())
                    description = st.text_area("Job Description")
                    voice_desc = voice_input_section("Experience Description via Voice", "voice_exp_desc")

                    if st.form_submit_button("Add Experience"):
                        full_desc = description + "\n" + voice_desc if voice_desc else description
                        st.session_state.resume_data["experience"].append({
                            "job_title": job_title,
                            "company": company,
                            "start_date": str(start_date),
                            "end_date": str(end_date),
                            "description": full_desc,
                        })
                        st.success("✅ Experience Added")
            with col_right:
                st.markdown("💡 Tip: Use short bullet points for impact.")

        # ==== EDUCATION TAB ====
        with tabs[2]:
            st.subheader("🎓 Education")
            col_left, col_right = st.columns([2, 1])
            with col_left:
                with st.form("education_form"):
                    degree = st.text_input("Degree")
                    school = st.text_input("School")
                    year = st.text_input("Year")
                    voice_education = voice_input_section("Education via Voice", "voice_education")

                    if st.form_submit_button("Add Education"):
                        full_degree = f"{degree} ({voice_education})" if voice_education else degree
                        st.session_state.resume_data["education"].append({
                            "degree": full_degree,
                            "school": school,
                            "year": year
                        })
                        st.success("✅ Education Added")
            with col_right:
                st.markdown("🎯 Include honors, GPA, or key coursework.")

        # ==== SKILLS TAB ====
        with tabs[3]:
            st.subheader("🛠️ Skills")
            col_left, col_right = st.columns([2, 1])
            with col_left:
                with st.form("skills_form"):
                    skills = st.text_area("List your skills")
                    voice_skills = voice_input_section("Skills via Voice", "voice_skills")

                    if st.form_submit_button("Save Skills"):
                        full_skills = skills + "\n" + voice_skills if voice_skills else skills
                        st.session_state.resume_data["skills"] = full_skills
                        st.success("✅ Skills Saved")
            with col_right:
                st.markdown("📌 Group skills by category.")

        # ==== CERTIFICATIONS TAB ====
        with tabs[4]:
            st.subheader("📜 Certifications")
            col_left, col_right = st.columns([2, 1])
            with col_left:
                with st.form("certifications_form"):
                    certifications = st.text_area("List your certifications")
                    voice_certifications = voice_input_section("Certifications via Voice", "voice_certifications")

                    if st.form_submit_button("Save Certifications"):
                        full_certifications = certifications + "\n" + voice_certifications if voice_certifications else certifications
                        st.session_state.resume_data["certifications"] = full_certifications
                        st.success("✅ Certifications Saved")
            with col_right:
                st.markdown("🏆 Include only relevant ones.")

        # ==== PREVIEW TAB ====
        with tabs[5]:
            st.subheader("📑 Resume Preview")
            
            st.markdown("""
            <style>
            .photo-box {
                    background-color: #f0f8ff;
                    border: 2px solid #007acc;
                    border-radius: 12px;
                    padding: 1rem;
                    text-align: center;
                    box-shadow: 0px 2px 6px rgba(0,0,0,0.1);
                }
                .photo-box img {
                    border-radius: 10px;
                }
                </style>
            """, unsafe_allow_html=True)
                    
                
            resume = st.session_state.resume_data

            col1, col2 = st.columns(2)

            with col1:
                
                if resume.get("personal_info"):
                    st.markdown("### 👤 Personal Information")
                    pi = resume["personal_info"]
                    st.write(f"**Name:** {pi.get('name','')}")
                    st.write(f"**Email:** {pi.get('email','')}")
                    st.write(f"**Phone:** {pi.get('phone','')}")
                    st.write(f"**Summary:** {pi.get('summary','')}")

                if resume.get("education"):
                    st.markdown("### 🎓 Education")
                    for edu in resume["education"]:
                        st.write(f"**{edu.get('degree','')}**, {edu.get('school','')} ({edu.get('year','')})")

            with col2:
                
                photo_data = resume.get("photo") or st.session_state.get("profile_photo")
                if photo_data:
                    st.markdown('<div class="photo-box">', unsafe_allow_html=True)
                    st.image(photo_data, width=200)
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.write("No photo uploaded")
                if resume.get("skills"):
                    st.markdown("### 🛠 Skills")
                    st.write(resume["skills"])

                if resume.get("experience"):
                    st.markdown("### 💼 Experience")
                    for exp in resume["experience"]:
                        st.write(f"**{exp.get('job_title','')}** at {exp.get('company','')}")
                        st.write(f"{exp.get('start_date','')} – {exp.get('end_date','')}")
                        st.write(exp.get("description",""))

                if resume.get("certifications"):
                    st.markdown("### 📜 Certifications")
                    st.write(resume["certifications"])
        
        
        create_prompt_template = (
            "You are a professional resume writer. Based on the description below, generate a structured and detailed resume with proper sections like Summary, Skills, Experience, Projects, and Education.\n\n"
            "Description:\n{text}\n\n"
            "Return in resume format as plain text."
        )
        
                    

        # Ensure edit mode state exists
        if "edit_mode" not in st.session_state:
            st.session_state.edit_mode = False

        # Generate Resume
        if st.button("✨ Generate Resume", key="generate_resume_btn"):
            combined_text = build_user_input(st.session_state.resume_data)
            
            if combined_text.strip():
                with st.spinner("Generating resume..."):
                    resume = summarize_with_gemini(combined_text, create_prompt_template)
                if resume:
                    st.success("✅ Resume Generated")
                    st.session_state.editable_resume = resume
                    st.session_state.edit_mode = False  # Start in view mode

                    # Save initial unedited docx path for download
                    docx_path = save_as_docx(resume, "generated_resume.docx")
                    st.session_state.docx_path = docx_path
                else:
                    st.error("Resume generation failed.")
            else:
                st.warning("Please fill in the form sections or provide input.")

        # Show Resume if available
        if "editable_resume" in st.session_state and st.session_state.editable_resume:
            if st.session_state.edit_mode:
                # --- Edit Mode ---
                edited_text = st.text_area(
                    "Edit Your Resume",
                    value=st.session_state.editable_resume,
                    height=500,
                    key="resume_editor"
                )
                if st.button("💾 Save Resume", key="save_resume_btn"):
                    st.session_state.editable_resume = edited_text
                    st.session_state.edit_mode = False

                    # Save edited version to docx
                    docx_path = save_as_docx(edited_text, "edited_resume.docx")
                    st.session_state.docx_path = docx_path

                    st.success("✅ Resume updated.")
                    st.rerun()  # Switch back to view mode instantly

            else:
                # --- View Mode ---
                st.markdown(markdown_to_html(st.session_state.editable_resume), unsafe_allow_html=True)

                col1, col2 = st.columns([1, 3])
                with col1:
                    if st.button("✏️ Edit Resume", key="edit_resume_btn"):
                        st.session_state.edit_mode = True
                        st.rerun()  # Switch to edit mode instantly

                st.markdown("<br>", unsafe_allow_html=True)  # Proper spacing before download

                # --- Download Button ---
                if "docx_path" in st.session_state:
                    st.download_button(
                        "📄 Download as DOCX",
                        data=open(st.session_state.docx_path, "rb"),
                        file_name="resume.docx"
                    )
             
            
    with tab2:
        st.subheader("Upload your existing resume and get it improved and more impactful")
        file = st.file_uploader("Upload Resume (PDF, DOCX, TXT, Image)", type=["pdf", "docx", "txt", "jpg", "jpeg", "png"])

        improve_prompt_template = (
            "You are an expert resume editor. Improve and rewrite the resume below to make it more impactful, clean, and suitable for top jobs.\n\n"
            "Resume:\n{text}\n\n"
            "Return the improved resume as plain text."
        )

        if file and st.button("🚀 Improve Resume"):
            with st.spinner("Extracting and improving resume..."):
                raw_text = extract_text_from_file(file)
                if raw_text:
                    improved = summarize_with_gemini(raw_text, improve_prompt_template)
                    if improved:
                        st.success("✅ Improved Resume")
                        st.markdown(markdown_to_html(improved), unsafe_allow_html=True)

                        
                        docx_path = save_as_docx(improved, "improved_resume.docx")
                        # pdf_path = save_as_pdf(improved, "improved_resume.pdf")

                        st.download_button("📄 Download as DOCX", data=open(docx_path, "rb"), file_name="improved_resume.docx")
                        # st.download_button("📄 Download as PDF", data=open(pdf_path, "rb"), file_name="improved_resume.pdf")
                        # st.download_button("📄 Download as TXT", data=improved, file_name="improved_resume.txt")

                    else:
                        st.error("Improvement failed.")
                else:
                    st.error("Could not extract text from the uploaded file.")

if __name__ == "__main__":
    main()
